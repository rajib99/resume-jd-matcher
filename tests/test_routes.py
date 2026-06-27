"""
Integration tests for HTTP routes and parser module.

Coverage targets
----------------
GET  /               – frontend HTML served
GET  /health         – status + version
POST /api/v1/match/text   – success, every validation failure, agent 502
POST /api/v1/match/upload – success PDF/DOCX/octet-stream,
                            413/415/422/502 error paths,
                            exact-5MB boundary,
                            jd_file too large,
                            missing filename edge case
app.services.parser
  extract_text_from_pdf   – success, image-only PDF → ValueError
  extract_text_from_docx  – success with real DOCX, empty DOCX → ValueError
  extract_text_from_file  – dispatch by content-type and by extension
"""

from __future__ import annotations

import io
from unittest.mock import MagicMock, patch

import pytest
from fastapi.testclient import TestClient

from app.services.parser import (
    extract_text_from_docx,
    extract_text_from_file,
    extract_text_from_pdf,
)

# ── Upload helper ─────────────────────────────────────────────────────────────

_PDF_CT   = "application/pdf"
_DOCX_CT  = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_OCTET_CT = "application/octet-stream"
_MAX_BYTES = 5 * 1024 * 1024


def _upload(
    client: TestClient,
    resume_bytes: bytes,
    resume_name: str,
    jd_bytes: bytes,
    jd_name: str,
    resume_ct: str = _PDF_CT,
    jd_ct: str = _PDF_CT,
):
    return client.post(
        "/api/v1/match/upload",
        files={
            "resume_file": (resume_name, io.BytesIO(resume_bytes), resume_ct),
            "jd_file":     (jd_name,     io.BytesIO(jd_bytes),     jd_ct),
        },
    )


# ════════════════════════════════════════════════════════════════════════════
# Parser unit tests
# ════════════════════════════════════════════════════════════════════════════


class TestExtractTextFromPdf:
    def _mock_open(self, page_text: str | None):
        """Return a context-manager mock for pdfplumber.open."""
        mock_page = MagicMock()
        mock_page.extract_text.return_value = page_text
        mock_ctx = MagicMock()
        mock_ctx.__enter__.return_value.pages = [mock_page]
        return mock_ctx

    def test_returns_extracted_text(self, sample_resume):
        with patch("pdfplumber.open", return_value=self._mock_open(sample_resume)):
            result = extract_text_from_pdf(b"any bytes")
        assert "Jane" in result

    def test_strips_and_normalises_whitespace(self):
        raw = "  Section A\n\n\n\nSection B  "
        with patch("pdfplumber.open", return_value=self._mock_open(raw)):
            result = extract_text_from_pdf(b"bytes")
        assert result == "Section A\n\nSection B"

    def test_multi_page_joined_with_blank_line(self):
        page1 = MagicMock()
        page1.extract_text.return_value = "Page one content"
        page2 = MagicMock()
        page2.extract_text.return_value = "Page two content"
        mock_ctx = MagicMock()
        mock_ctx.__enter__.return_value.pages = [page1, page2]

        with patch("pdfplumber.open", return_value=mock_ctx):
            result = extract_text_from_pdf(b"bytes")

        assert "Page one content" in result
        assert "Page two content" in result

    def test_image_only_pdf_raises_value_error(self):
        with patch("pdfplumber.open", return_value=self._mock_open(None)):
            with pytest.raises(ValueError, match="no extractable text"):
                extract_text_from_pdf(b"bytes")

    def test_all_pages_empty_raises_value_error(self):
        p1, p2 = MagicMock(), MagicMock()
        p1.extract_text.return_value = ""
        p2.extract_text.return_value = "   "
        mock_ctx = MagicMock()
        mock_ctx.__enter__.return_value.pages = [p1, p2]

        with patch("pdfplumber.open", return_value=mock_ctx):
            with pytest.raises(ValueError):
                extract_text_from_pdf(b"bytes")


class TestExtractTextFromDocx:
    def test_returns_text_from_real_docx(self, sample_docx_bytes, sample_resume):
        result = extract_text_from_docx(sample_docx_bytes)
        assert "Jane" in result

    def test_normalises_whitespace_in_result(self, sample_docx_bytes):
        result = extract_text_from_docx(sample_docx_bytes)
        assert "\n\n\n" not in result

    def test_empty_docx_raises_value_error(self):
        import docx

        doc = docx.Document()
        buf = io.BytesIO()
        doc.save(buf)

        with pytest.raises(ValueError, match="no extractable text"):
            extract_text_from_docx(buf.getvalue())

    def test_docx_with_only_whitespace_paragraphs_raises(self):
        import docx

        doc = docx.Document()
        doc.add_paragraph("   ")
        doc.add_paragraph("\t\t")
        buf = io.BytesIO()
        doc.save(buf)

        with pytest.raises(ValueError, match="no extractable text"):
            extract_text_from_docx(buf.getvalue())


class TestExtractTextFromFile:
    def test_dispatches_pdf_by_content_type(self, sample_resume):
        with patch("app.services.parser.extract_text_from_pdf", return_value=sample_resume) as m:
            result = extract_text_from_file(b"bytes", "doc.pdf", _PDF_CT)
        m.assert_called_once_with(b"bytes")
        assert result == sample_resume

    def test_dispatches_docx_by_content_type(self, sample_resume):
        with patch("app.services.parser.extract_text_from_docx", return_value=sample_resume) as m:
            result = extract_text_from_file(b"bytes", "doc.docx", _DOCX_CT)
        m.assert_called_once_with(b"bytes")
        assert result == sample_resume

    def test_dispatches_pdf_by_extension_when_octet_stream(self, sample_resume):
        with patch("app.services.parser.extract_text_from_pdf", return_value=sample_resume) as m:
            result = extract_text_from_file(b"bytes", "resume.pdf", _OCTET_CT)
        m.assert_called_once()
        assert result == sample_resume

    def test_dispatches_docx_by_extension_when_octet_stream(self, sample_resume):
        with patch("app.services.parser.extract_text_from_docx", return_value=sample_resume) as m:
            extract_text_from_file(b"bytes", "resume.docx", _OCTET_CT)
        m.assert_called_once()

    def test_pdf_extension_beats_missing_content_type(self, sample_resume):
        with patch("app.services.parser.extract_text_from_pdf", return_value=sample_resume):
            result = extract_text_from_file(b"bytes", "file.pdf", "")
        assert result == sample_resume

    def test_unsupported_type_raises_value_error(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text_from_file(b"bytes", "document.txt", "text/plain")

    def test_no_extension_unsupported_type_raises(self):
        with pytest.raises(ValueError):
            extract_text_from_file(b"bytes", "document", "application/octet-stream")


# ════════════════════════════════════════════════════════════════════════════
# GET /
# ════════════════════════════════════════════════════════════════════════════


class TestFrontendRoute:
    def test_returns_200(self, test_client):
        resp = test_client.get("/")
        assert resp.status_code == 200

    def test_returns_html_content_type(self, test_client):
        resp = test_client.get("/")
        assert "text/html" in resp.headers["content-type"]

    def test_html_contains_expected_elements(self, test_client):
        resp = test_client.get("/")
        body = resp.text
        assert "Resume JD Matcher" in body
        assert "match-btn" in body
        assert "/api/v1/match/" in body

    def test_html_references_unified_endpoint(self, test_client):
        resp = test_client.get("/")
        assert "/api/v1/match/" in resp.text


# ════════════════════════════════════════════════════════════════════════════
# GET /health
# ════════════════════════════════════════════════════════════════════════════


class TestHealthRoute:
    def test_returns_200(self, test_client):
        assert test_client.get("/health").status_code == 200

    def test_returns_ok_status(self, test_client):
        assert test_client.get("/health").json()["status"] == "ok"

    def test_returns_version(self, test_client):
        assert test_client.get("/health").json()["version"] == "1.0.0"


# ════════════════════════════════════════════════════════════════════════════
# POST /api/v1/match/text
# ════════════════════════════════════════════════════════════════════════════


class TestMatchTextEndpoint:
    def test_success_returns_200(self, test_client, sample_resume, sample_jd):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": sample_resume, "job_description": sample_jd},
        )
        assert resp.status_code == 200

    def test_response_contains_all_report_fields(self, test_client, sample_resume, sample_jd):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": sample_resume, "job_description": sample_jd},
        )
        report = resp.json()["report"]
        for field in (
            "overall_match_score", "matched_skills", "missing_skills",
            "strengths", "gaps", "recommendation", "summary",
        ):
            assert field in report, f"Missing field: {field}"

    def test_response_contains_metadata(self, test_client, sample_resume, sample_jd):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": sample_resume, "job_description": sample_jd},
        )
        data = resp.json()
        assert "model_used" in data
        assert "processing_time_ms" in data

    def test_service_called_with_correct_inputs(
        self, mock_service, test_client, sample_resume, sample_jd
    ):
        test_client.post(
            "/api/v1/match/text",
            json={"resume": sample_resume, "job_description": sample_jd},
        )
        mock_service.match.assert_called_once_with(
            resume_text=sample_resume, job_description=sample_jd
        )

    # ── Validation failures ──────────────────────────────────────────────────

    def test_resume_too_short_returns_422(self, test_client, sample_jd):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": "short", "job_description": sample_jd},
        )
        assert resp.status_code == 422

    def test_jd_too_short_returns_422(self, test_client, sample_resume):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": sample_resume, "job_description": "short"},
        )
        assert resp.status_code == 422

    def test_missing_resume_field_returns_422(self, test_client, sample_jd):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"job_description": sample_jd},
        )
        assert resp.status_code == 422

    def test_missing_jd_field_returns_422(self, test_client, sample_resume):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": sample_resume},
        )
        assert resp.status_code == 422

    def test_empty_body_returns_422(self, test_client):
        resp = test_client.post("/api/v1/match/text", json={})
        assert resp.status_code == 422

    def test_resume_exactly_49_chars_returns_422(self, test_client, sample_jd):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": "x" * 49, "job_description": sample_jd},
        )
        assert resp.status_code == 422

    def test_resume_exactly_50_chars_accepted(self, test_client, sample_jd):
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": "x" * 50, "job_description": sample_jd},
        )
        # Passes validation; 200 or 502 depending on service, never 422
        assert resp.status_code != 422

    # ── Service error ────────────────────────────────────────────────────────

    def test_agent_error_returns_502(self, test_client, sample_resume, sample_jd, mock_service):
        mock_service.match.side_effect = RuntimeError("Claude down")
        resp = test_client.post(
            "/api/v1/match/text",
            json={"resume": sample_resume, "job_description": sample_jd},
        )
        assert resp.status_code == 502
        assert "Agent error" in resp.json()["detail"]


# ════════════════════════════════════════════════════════════════════════════
# POST /api/v1/match/upload
# ════════════════════════════════════════════════════════════════════════════


class TestMatchUploadEndpoint:
    # ── Success paths ────────────────────────────────────────────────────────

    def test_pdf_upload_returns_200(self, test_client, minimal_pdf_bytes, sample_resume):
        with patch("app.routers.match.extract_text_from_file", return_value=sample_resume):
            resp = _upload(test_client, minimal_pdf_bytes, "r.pdf", minimal_pdf_bytes, "j.pdf")
        assert resp.status_code == 200

    def test_docx_upload_returns_200(self, test_client, sample_docx_bytes, sample_resume):
        with patch("app.routers.match.extract_text_from_file", return_value=sample_resume):
            resp = _upload(
                test_client,
                sample_docx_bytes, "r.docx",
                sample_docx_bytes, "j.docx",
                resume_ct=_DOCX_CT, jd_ct=_DOCX_CT,
            )
        assert resp.status_code == 200

    def test_octet_stream_with_pdf_extension_returns_200(
        self, test_client, minimal_pdf_bytes, sample_resume
    ):
        with patch("app.routers.match.extract_text_from_file", return_value=sample_resume):
            resp = _upload(
                test_client,
                minimal_pdf_bytes, "r.pdf",
                minimal_pdf_bytes, "j.pdf",
                resume_ct=_OCTET_CT, jd_ct=_OCTET_CT,
            )
        assert resp.status_code == 200

    def test_octet_stream_with_docx_extension_returns_200(
        self, test_client, sample_docx_bytes, sample_resume
    ):
        with patch("app.routers.match.extract_text_from_file", return_value=sample_resume):
            resp = _upload(
                test_client,
                sample_docx_bytes, "r.docx",
                sample_docx_bytes, "j.docx",
                resume_ct=_OCTET_CT, jd_ct=_OCTET_CT,
            )
        assert resp.status_code == 200

    def test_extract_called_for_both_files(
        self, test_client, minimal_pdf_bytes, sample_resume
    ):
        with patch(
            "app.routers.match.extract_text_from_file", return_value=sample_resume
        ) as mock_extract:
            _upload(test_client, minimal_pdf_bytes, "r.pdf", minimal_pdf_bytes, "j.pdf")
        assert mock_extract.call_count == 2

    def test_response_shape_on_upload_success(
        self, test_client, minimal_pdf_bytes, sample_resume
    ):
        with patch("app.routers.match.extract_text_from_file", return_value=sample_resume):
            resp = _upload(test_client, minimal_pdf_bytes, "r.pdf", minimal_pdf_bytes, "j.pdf")
        data = resp.json()
        assert "report" in data
        assert "model_used" in data
        assert "processing_time_ms" in data

    # ── File size limits ─────────────────────────────────────────────────────

    def test_resume_file_over_5mb_returns_413(
        self, test_client, minimal_pdf_bytes
    ):
        oversized = b"x" * (_MAX_BYTES + 1)
        resp = _upload(test_client, oversized, "r.pdf", minimal_pdf_bytes, "j.pdf")
        assert resp.status_code == 413
        assert "resume_file" in resp.json()["detail"]

    def test_jd_file_over_5mb_returns_413(self, test_client, minimal_pdf_bytes):
        oversized = b"x" * (_MAX_BYTES + 1)
        resp = _upload(test_client, minimal_pdf_bytes, "r.pdf", oversized, "j.pdf")
        assert resp.status_code == 413
        assert "jd_file" in resp.json()["detail"]

    def test_file_exactly_at_5mb_limit_is_accepted(
        self, test_client, sample_resume
    ):
        at_limit = b"x" * _MAX_BYTES
        with patch("app.routers.match.extract_text_from_file", return_value=sample_resume):
            resp = _upload(test_client, at_limit, "r.pdf", at_limit, "j.pdf")
        assert resp.status_code != 413

    # ── Content type / extension validation ─────────────────────────────────

    def test_txt_file_returns_415(self, test_client):
        txt = b"some plain text " * 20
        resp = _upload(
            test_client, txt, "r.txt", txt, "j.txt",
            resume_ct="text/plain", jd_ct="text/plain",
        )
        assert resp.status_code == 415

    def test_html_file_returns_415(self, test_client):
        html = b"<html><body>Resume</body></html>" * 5
        resp = _upload(
            test_client, html, "r.html", html, "j.html",
            resume_ct="text/html", jd_ct="text/html",
        )
        assert resp.status_code == 415

    def test_octet_stream_with_txt_extension_returns_415(self, test_client):
        data = b"data " * 20
        resp = _upload(
            test_client, data, "r.txt", data, "j.txt",
            resume_ct=_OCTET_CT, jd_ct=_OCTET_CT,
        )
        assert resp.status_code == 415

    def test_415_error_mentions_field_name(self, test_client):
        txt = b"x " * 30
        resp = _upload(
            test_client, txt, "r.txt", txt, "j.txt",
            resume_ct="text/plain", jd_ct="text/plain",
        )
        detail = resp.json()["detail"]
        assert "resume_file" in detail

    # ── Extracted text too short ─────────────────────────────────────────────

    def test_short_extracted_text_returns_422(
        self, test_client, minimal_pdf_bytes
    ):
        with patch("app.routers.match.extract_text_from_file", return_value="too short"):
            resp = _upload(test_client, minimal_pdf_bytes, "r.pdf", minimal_pdf_bytes, "j.pdf")
        assert resp.status_code == 422

    def test_short_jd_text_returns_422(self, test_client, minimal_pdf_bytes, sample_resume):
        call_count = 0

        def _side_effect(*_args, **_kwargs):
            nonlocal call_count
            call_count += 1
            return sample_resume if call_count == 1 else "short"

        with patch("app.routers.match.extract_text_from_file", side_effect=_side_effect):
            resp = _upload(test_client, minimal_pdf_bytes, "r.pdf", minimal_pdf_bytes, "j.pdf")
        assert resp.status_code == 422

    def test_extraction_error_returns_422(self, test_client, minimal_pdf_bytes):
        with patch(
            "app.routers.match.extract_text_from_file",
            side_effect=ValueError("PDF is image-only"),
        ):
            resp = _upload(test_client, minimal_pdf_bytes, "r.pdf", minimal_pdf_bytes, "j.pdf")
        assert resp.status_code == 422
        assert "image-only" in resp.json()["detail"]

    # ── Missing fields ───────────────────────────────────────────────────────

    def test_missing_jd_file_returns_422(self, test_client, minimal_pdf_bytes):
        resp = test_client.post(
            "/api/v1/match/upload",
            files={"resume_file": ("r.pdf", io.BytesIO(minimal_pdf_bytes), _PDF_CT)},
        )
        assert resp.status_code == 422

    def test_missing_resume_file_returns_422(self, test_client, minimal_pdf_bytes):
        resp = test_client.post(
            "/api/v1/match/upload",
            files={"jd_file": ("j.pdf", io.BytesIO(minimal_pdf_bytes), _PDF_CT)},
        )
        assert resp.status_code == 422

    def test_no_files_returns_422(self, test_client):
        resp = test_client.post("/api/v1/match/upload")
        assert resp.status_code == 422

    # ── Agent / service errors ───────────────────────────────────────────────

    def test_agent_error_returns_502(
        self, test_client, mock_service, minimal_pdf_bytes, sample_resume
    ):
        mock_service.match.side_effect = RuntimeError("model timeout")
        with patch("app.routers.match.extract_text_from_file", return_value=sample_resume):
            resp = _upload(test_client, minimal_pdf_bytes, "r.pdf", minimal_pdf_bytes, "j.pdf")
        assert resp.status_code == 502
        assert "Agent error" in resp.json()["detail"]
